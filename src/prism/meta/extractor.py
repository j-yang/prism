"""Mock Shell Document Extractor.

Parses docx/xlsx mock shells into structured JSON for spec generation.
"""

import json
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document


@dataclass
class Deliverable:
    """Represents a TLF deliverable from mock shell."""

    deliverable_id: str
    deliverable_type: str  # 'table', 'listing', 'figure'
    title: str
    population: str = ""
    schema: str = ""  # 'baseline', 'longitudinal', 'occurrence'
    columns: list = field(default_factory=list)
    rows: list = field(default_factory=list)
    footnotes: dict = field(default_factory=dict)
    programming_notes: list = field(default_factory=list)
    table_data: list = field(default_factory=list)


@dataclass
class MockShellContext:
    """Structured context extracted from mock shell."""

    study_title: str = ""
    protocol_no: str = ""
    deliverables: list = field(default_factory=list)
    sections: dict = field(default_factory=dict)
    all_footnotes: dict = field(default_factory=dict)
    all_programming_notes: list = field(default_factory=list)


class MockShellExtractor:
    """Extract structured context from mock shell documents."""

    DELIVERABLE_PATTERN = re.compile(
        r"^(Table|Listing|Figure)\s+(\d+\.\d+\.\d+(?:\.\d+)?)\s+(.+)$", re.IGNORECASE
    )

    POPULATION_PATTERNS = [
        r"\(Safety Set\)",
        r"\(Full Analysis Set\)",
        r"\(FAS\)",
        r"\(SS\)",
        r"<<IIM Cohort>>",
        r"<<SSc Cohort>>",
    ]

    FOOTNOTE_PATTERN = re.compile(r"\[([a-z])\]")

    def __init__(self):
        self.context = MockShellContext()

    def extract(self, doc_path: str) -> MockShellContext:
        """Extract context from mock shell document."""
        path = Path(doc_path)

        if path.suffix.lower() == ".docx":
            self._extract_from_docx(path)
        elif path.suffix.lower() == ".xlsx":
            self._extract_from_xlsx(path)
        else:
            raise ValueError(f"Unsupported file format: {path.suffix}")

        return self.context

    def _extract_from_docx(self, path: Path):
        """Extract from Word document."""
        doc = Document(str(path))

        self._extract_header_info(doc)

        body = doc.element.body
        children = list(body.iterchildren())

        deliverables_by_element = {}
        current_section = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            section_match = re.match(r"^(\d+\.\d+)\s+(.+)$", text)
            if section_match:
                current_section = section_match.group(1)
                self.context.sections[current_section] = {
                    "title": section_match.group(2),
                    "deliverables": [],
                }
                continue

            deliv_match = self.DELIVERABLE_PATTERN.match(text)
            if deliv_match:
                deliverable = Deliverable(
                    deliverable_id=deliv_match.group(2),
                    deliverable_type=deliv_match.group(1).lower(),
                    title=deliv_match.group(3).strip(),
                )

                population = self._extract_population(text)
                if population:
                    deliverable.population = population

                if current_section:
                    self.context.sections[current_section]["deliverables"].append(
                        deliverable.deliverable_id
                    )

                deliverables_by_element[para._element] = deliverable

        self._extract_tables_by_element(doc, children, deliverables_by_element)

        seen_keys = {}
        for para_elem, deliverable in deliverables_by_element.items():
            key = f"{deliverable.deliverable_type}_{deliverable.deliverable_id}"
            if key not in seen_keys:
                seen_keys[key] = deliverable
            else:
                existing = seen_keys[key]
                if deliverable.table_data and not existing.table_data:
                    seen_keys[key] = deliverable

        for deliverable in seen_keys.values():
            self.context.deliverables.append(deliverable)

    def _extract_tables_by_element(
        self, doc: Document, children: list, deliverables_by_element: dict
    ):
        """Extract tables and associate with deliverables based on element position."""
        if not deliverables_by_element:
            return

        para_to_deliverable = {}
        for para_elem, deliverable in deliverables_by_element.items():
            para_to_deliverable[id(para_elem)] = deliverable

        deliverable_elements = set(deliverables_by_element.keys())

        for table in doc.tables:
            table_elem = table._element

            try:
                table_idx = children.index(table_elem)
            except ValueError:
                continue

            best_deliverable = None
            best_distance = float("inf")

            for i in range(table_idx - 1, max(0, table_idx - 50), -1):
                child = children[i]
                if child in deliverable_elements:
                    deliverable = deliverables_by_element[child]
                    if not deliverable.table_data:
                        best_deliverable = deliverable
                        break

            if best_deliverable:
                rows_data = []
                for row in table.rows:
                    row_text = [
                        cell.text.strip().replace("\n", " | ") for cell in row.cells
                    ]
                    rows_data.append(row_text)

                if rows_data and any(any(cell for cell in row) for row in rows_data):
                    best_deliverable.table_data = rows_data
                    if len(rows_data) > 0:
                        best_deliverable.columns = rows_data[0]
                    if len(rows_data) > 1:
                        best_deliverable.rows = [
                            r[0] for r in rows_data[1:] if r and r[0]
                        ]

                    best_deliverable.schema = self._infer_schema(best_deliverable)

    def _extract_from_xlsx(self, path: Path):
        """Extract from Excel document."""
        import openpyxl

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]

            deliv_match = self.DELIVERABLE_PATTERN.match(sheet_name)
            if deliv_match:
                deliverable = Deliverable(
                    deliverable_id=deliv_match.group(2),
                    deliverable_type=deliv_match.group(1).lower(),
                    title=deliv_match.group(3).strip(),
                )

                rows_data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        rows_data.append([str(cell) if cell else "" for cell in row])

                if rows_data:
                    deliverable.table_data = rows_data
                    if len(rows_data) > 0:
                        deliverable.columns = rows_data[0]
                    if len(rows_data) > 1:
                        deliverable.rows = [r[0] for r in rows_data[1:] if r]

                self.context.deliverables.append(deliverable)

        wb.close()

    def _extract_header_info(self, doc: Document):
        """Extract study header information."""
        for para in doc.paragraphs[:20]:
            text = para.text.strip()

            if "Protocol No" in text or "Protocol No." in text:
                match = re.search(r"Protocol No\.?:?\s*(.+)", text)
                if match:
                    self.context.protocol_no = match.group(1).strip()

            if "Open-Label Study" in text or "Study to Assess" in text:
                self.context.study_title = text

    def _extract_population(self, text: str) -> str:
        """Extract population from text."""
        populations = []

        if re.search(r"\(Safety Set\)", text, re.IGNORECASE):
            populations.append("Safety Set")
        if re.search(r"\(Full Analysis Set\)", text, re.IGNORECASE):
            populations.append("Full Analysis Set")
        if re.search(r"<<IIM Cohort>>", text):
            populations.append("IIM Cohort")
        if re.search(r"<<SSc Cohort>>", text):
            populations.append("SSc Cohort")

        return ", ".join(populations)

    def _infer_schema(self, deliverable: Deliverable) -> str:
        """Infer schema type from deliverable content."""
        text_to_check = " ".join(
            [
                deliverable.title.lower(),
                " ".join(deliverable.rows).lower(),
                " ".join(deliverable.columns).lower(),
            ]
        )

        if any(
            kw in text_to_check
            for kw in [
                "adverse event",
                "ae",
                "sae",
                "crs",
                "icans",
                "medication",
                "conmed",
            ]
        ):
            return "occurrence"
        elif any(
            kw in text_to_check
            for kw in ["visit", "over time", "longitudinal", "by visit"]
        ):
            return "longitudinal"
        else:
            return "baseline"

    def to_json(self) -> str:
        """Export context as JSON string."""
        return json.dumps(asdict(self.context), indent=2, ensure_ascii=False)

    def save_json(self, output_path: str):
        """Save context to JSON file."""
        Path(output_path).write_text(self.to_json(), encoding="utf-8")


def extract_mock_shell(
    doc_path: str, output_json: Optional[str] = None
) -> MockShellContext:
    """Convenience function to extract mock shell.

    Args:
        doc_path: Path to mock shell document (docx/xlsx)
        output_json: Optional path to save JSON output

    Returns:
        MockShellContext with extracted data
    """
    extractor = MockShellExtractor()
    context = extractor.extract(doc_path)

    if output_json:
        extractor.save_json(output_json)

    return context
